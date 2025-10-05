from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import openpyxl
from datetime import datetime, timedelta  # 添加 timedelta 导入




def excel_to_word(excel_path, word_path):
    # 读取Excel文件
    wb = openpyxl.load_workbook(excel_path)

    # 创建Word文档
    doc = Document()

    # 设置文档标题
    title = doc.add_paragraph()
    title_run = title.add_run("作品汇总")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 总序号
    global_index = 1

    # 遍历所有sheet（假设处理前两个sheet）
    for sheet_name in wb.sheetnames[:2]:
        sheet = wb[sheet_name]

        # 添加sheet标题（可选）
        sheet_title = doc.add_paragraph()
        sheet_title_run = sheet_title.add_run(f"\n--- {sheet_name} ---")
        sheet_title_run.font.size = Pt(14)
        sheet_title_run.font.bold = True

        # 处理当前sheet的数据
        for row in sheet.iter_rows(min_row=2, values_only=True):
            name = row[0]  # A列
            time = row[1]
            ip = row[2]
            company = row[3]  # C列

            # 添加序号
            doc.add_paragraph(f"{global_index}.")

            # 添加作品名称
            name_para = doc.add_paragraph()
            name_para.add_run("作品名称：").bold = True
            name_para.add_run(str(name) if name is not None else "")


            # 开播时间 - 特殊处理日期时间类型
            time_text = ""
            if time is not None:
                if isinstance(time, datetime):
                    time_text = time.strftime('%Y-%m-%d %H:%M:%S')
                elif isinstance(time, (int, float)):
                    # 处理Excel日期序列化格式
                    time_obj = datetime(1899, 12, 30) + timedelta(days=time)
                    time_text = time_obj.strftime('%Y-%m-%d')
                else:
                    time_text = str(time)
            time_para = doc.add_paragraph()
            time_para.add_run("播出时间：").bold = True
            time_para.add_run(time_text)

            # 添加所属公司
            company_para = doc.add_paragraph()
            company_para.add_run("作品所属公司：").bold = True
            company_para.add_run(str(company) if company is not None else "")

            # 添加IP字段
            ip_para = doc.add_paragraph()
            ip_para.add_run("常见ip：").bold = True
            ip_para.add_run(str(ip) if ip is not None else "")

            # 添加空行分隔
            doc.add_paragraph()

            # 序号递增
            global_index += 1

    # 保存Word文档
    doc.save(word_path)
    print(f"Word文档已生成: {word_path}")

# 使用示例
if __name__ == "__main__":
    excel = "C:/Users/Clara/Desktop/动漫动画/数据收集/2024.11-2025.6动画电影.xlsx"
    word = "C:/Users/Clara/Desktop/动漫动画/数据收集/3作品汇总.docx"
    excel_to_word(excel, word)