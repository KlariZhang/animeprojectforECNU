from docx import Document

# 读取原始文档
doc = Document("整合指标8.23 - 副本.docx")

# 新建一个文档用于输出修改版本
new_doc = Document()

# 定义统一的处理方式描述
standard_method = "处理方式：所有指标均采用统计方法进行标准化（如 log(1+X) 转换后再进行 Min-Max 标准化，或直接采用 Z-score 标准化），不使用其他复杂方法。"
equal_weight_text = "权重：所有一级指标与二级指标均为等权重分配。"

# 遍历原文档内容，进行替换
for para in doc.paragraphs:
    text = para.text

    # 删除“权重×...”等复杂公式，替换为等权重描述
    if "权重" in text:
        text = equal_weight_text

    # 删除贝叶斯修正、加权公式、复杂方法
    if "贝叶斯" in text or "加权" in text or "公式" in text or "Score" in text:
        text = standard_method

    # 删除情感加权
    if "情感" in text:
        text = "评论情感分析：基于评论情绪标签计数，统一采用统计标准化方法。"

    # 数据处理方式部分统一替换
    if "处理方式" in text or "计算方式" in text:
        text = standard_method

    # 将修改后的段落写入新文档
    new_doc.add_paragraph(text)

# 保存新文档
output_path = "整合指标8.23_等权重统计化版.docx"
new_doc.save(output_path)


